from docx import Document

def parse_txt(text_bytes: bytes) -> dict:
    text = text_bytes.decode("utf-8", errors="ignore")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    # naive heuristics
    data = {"experience": [], "education": [], "skills": []}
    # First non-empty line as name guess
    if lines:
        data["name"] = lines[0]
    # Find sections by keywords
    sections = {"summary": [], "experience": [], "education": [], "skills": []}
    current = "summary"
    for ln in lines[1:]:
        low = ln.lower()
        if "experience" in low:
            current = "experience"; continue
        if "education" in low:
            current = "education"; continue
        if "skill" in low:
            current = "skills"; continue
        sections[current].append(ln)
    if sections["summary"]:
        data["summary"] = " ".join(sections["summary"][:4])
    if sections["skills"]:
        # split by commas
        toks = []
        for s in sections["skills"]:
            toks.extend([t.strip() for t in s.split(",") if t.strip()])
        data["skills"] = toks[:30]
    if sections["experience"]:
        # group every 4-5 lines into one job
        chunk, buf = [], []
        for s in sections["experience"]:
            if s.endswith(".") or "-" in s:
                buf.append(s)
                if len(buf) >= 5:
                    chunk.append(buf); buf = []
            else:
                buf.append(s)
        if buf: chunk.append(buf)
        for c in chunk:
            bullets = [x for x in c[1:] if x and not x.isupper()]
            data["experience"].append({
                "role": c[0] if c else "",
                "company": "",
                "dates": "",
                "location": "",
                "bullets": bullets[:6]
            })
    return data

def parse_docx(file_bytes: bytes) -> dict:
    # Save temp, open with python-docx
    import tempfile, os
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    doc = Document(tmp_path)
    os.unlink(tmp_path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    data = {"experience": [], "education": [], "skills": []}
    if paras:
        data["name"] = paras[0]
    # Section detection by headings and keywords
    sections = {"summary": [], "experience": [], "education": [], "skills": []}
    current = "summary"
    for p in doc.paragraphs[1:]:
        text = p.text.strip()
        if not text: continue
        style_name = (p.style.name or "").lower()
        tlow = text.lower()
        if any(k in tlow for k in ["experience"]):
            current = "experience"; continue
        if any(k in tlow for k in ["education"]):
            current = "education"; continue
        if any(k in tlow for k in ["skill"]):
            current = "skills"; continue
        sections[current].append(text)
    if sections["summary"]:
        data["summary"] = " ".join(sections["summary"][:4])
    if sections["skills"]:
        toks = []
        for s in sections["skills"]:
            toks.extend([t.strip() for t in s.split(",") if t.strip()])
        data["skills"] = toks[:30]
    if sections["experience"]:
        # Simple bullet grouping: blank lines or bullet prefixes
        job = {"role":"","company":"","dates":"","location":"","bullets":[]}
        for line in sections["experience"]:
            if line.startswith(("-", "•")):
                job["bullets"].append(line.lstrip("-• ").strip())
            elif any(sep in line for sep in ["—", " - "]):
                # treat as a header line "Role — Company"
                parts = [x.strip() for x in line.replace("—","-").split("-")]
                if job["role"] or job["bullets"]:
                    data["experience"].append(job); job = {"role":"","company":"","dates":"","location":"","bullets":[]}
                if parts:
                    job["role"] = parts[0]
                if len(parts) > 1:
                    job["company"] = parts[1]
            else:
                # free text appended to bullets
                if line:
                    job["bullets"].append(line)
        if job["role"] or job["bullets"]:
            data["experience"].append(job)
    return data
